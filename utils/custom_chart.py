from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# for chinese show
plt.rcParams['axes.unicode_minus'] = False
myfont = fm.FontProperties(fname='微軟正黑體_1.ttf')


def plt_bar(names, values, picture):
    plt.figure(figsize=(9, 3))
    plt.bar(names, values)  # bar
    plt.suptitle('柱状图', fontproperties=myfont)
    plt.savefig(picture)  # eps, jpeg, jpg, pdf, pgf, png, ps, raw, rgba, svg, svgz, tif, tiff 都可以
    # plt.show()


def plt_plot(names, values, picture):
    plt.figure(figsize=(9, 3))
    plt.plot(names, values)  # line
    plt.suptitle('折线图', fontproperties=myfont)
    plt.savefig(picture)  # eps, jpeg, jpg, pdf, pgf, png, ps, raw, rgba, svg, svgz, tif, tiff 都可以
    # plt.show()


def plt_pie(labels, sizes, picture):
    # Pie chart, where the slices will be ordered and plotted counter-clockwise:
    explode = (0, 0, 0)  # only "explode" the 2nd slice (i.e. 'Hogs')
    fig1, ax1 = plt.subplots()
    ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%', startangle=90)
    ax1.pie([1, 0, 0], radius=0.6, colors='w')
    ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    plt.suptitle('饼状图', fontproperties=myfont)
    plt.savefig(picture)  # eps, jpeg, jpg, pdf, pgf, png, ps, raw, rgba, svg, svgz, tif, tiff 都可以
    plt.show()


def demo():
    document = Document()
    # title
    document.add_heading('Document Title', 0)

    # paragraph
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True  # bold
    p.add_run(' and some ')
    p.add_run('italic.').italic = True  # italic

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    # make picture

    x_names = ['group_a', 'group_b', 'group_c']
    y_values = [1, 10, 100]
    plt_bar(x_names, y_values, "柱状图.png")
    plt_plot(x_names, y_values, "折线图.png")

    labels = 'Low,3', 'Medium,2', 'High,1', 'Logs'
    sizes = [3, 2, 1, 0]
    plt_pie(labels, sizes, "饼状图.png")

    document.add_picture('柱状图.png', width=Inches(6.25))
    document.add_picture('折线图.png', width=Inches(6.25))
    document.add_picture('饼状图.png', width=Inches(6.25))
    document.add_picture('散点图.png', width=Inches(6.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    # next pages
    document.add_page_break()

    document.save('demo.docx')


if __name__ == '__main__':
    main_labels = 'Low,3', 'Medium,2', 'High,1'
    sizes = [3, 2, 1]
    plt_pie(main_labels, sizes, "饼状图.png")