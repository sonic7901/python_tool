import os
import json
import pathlib
import re
import shutil
import pandas
from datetime import datetime, timedelta
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Mm, Inches
from docxtpl import DocxTemplate, InlineImage
import utils.custom_chrome
import custom_image
import custom_chart


# manual input
default_company = ''
default_date_before = ''
default_date_start = ''
default_date_end = ''
default_screenshot_url = ''


def set_company(input_company):
    global default_company
    default_company = input_company


def read_json(report_name):
    dict_data = ""
    try:
        if os.path.isfile(report_name):
            with open(report_name, 'r', encoding='utf-8') as read_file:
                dict_data = json.load(read_file)
    except Exception as ex:
        print('Exception:' + str(ex))
    return dict_data


def read_date_today():
    time_now = datetime.now()
    str_now = time_now.strftime("%Y/%m/%d")
    return str_now


def read_date_yesterday():
    time_yesterday = datetime.now() - timedelta(1)
    str_yesterday = time_yesterday.strftime("%Y/%m/%d")
    return str_yesterday


def read_screenshot(input_url, input_filename):
    main_path = pathlib.Path(__file__).parent.resolve()
    utils.custom_chrome.read_get_page_screenshot(input_url, str(main_path) + '\\', input_filename)


def write_variable(temp_document, replace_dict, input_fn):
    doc = DocxTemplate(temp_document)  # 选定模板
    doc.render(replace_dict)  # 渲染替换
    doc.save(input_fn)  # 保存


def remove_punctuation(text):
    try:
        punctuation = '!,;:?"\''
        text = re.sub(r'[{}]+'.format(punctuation), '', text)
        return text.strip().lower()
    except Exception as ex:
        print('Exception' + str(ex))
        return ''


def add_summary(input_docx, input_id, input_name, input_risk, input_filename, input_en):
    input_document = Document(input_docx)
    tables = input_document.tables
    table = tables[1]
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].paragraph_format.first_line_indent = Cm(0.45)
    row_cells[0].paragraphs[0].add_run(input_id)
    row_cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
    row_cells[1].paragraphs[0].add_run(input_name)
    if input_en:
        row_cells[2].paragraphs[0].paragraph_format.first_line_indent = Cm(0.45)
        row_cells[2].paragraphs[0].add_run(input_risk)
    else:
        row_cells[2].text = input_risk
    input_document.save(input_filename)
    return


def add_target(input_docx, input_name, input_url):
    input_document = Document(input_docx)
    tables = input_document.tables
    table = tables[0]
    row_cells = table.add_row().cells
    row_cells[0].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
    row_cells[0].paragraphs[0].add_run(input_name)
    row_cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
    row_cells[1].paragraphs[0].add_run(input_url)
    input_document.save(input_docx)
    return


def add_screenshot(input_docx, input_image, input_title):
    input_document = Document(input_docx)
    x = input_document.paragraphs
    paragraph = x[36]
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    temp_space = paragraph.add_run(input_title)
    temp_space.add_picture(input_image, width=Inches(7.5))
    input_document.save(input_docx)


def add_issue(input_docx,
              input_id,
              input_name,
              input_risk,
              input_type,
              input_target,
              input_range,
              input_detail,
              input_advice,
              input_info,
              input_fn,
              input_en):
    # create issue table
    input_document = Document(input_docx)
    if input_en:
        document_table = Document('template_table_issue_en.docx')
    else:
        document_table = Document('template_table_issue_zh.docx')
    temp_table = document_table.tables
    table_issue = temp_table[0]
    table_issue.rows[0].cells[1].paragraphs[0].add_run(input_id)
    table_issue.rows[1].cells[1].paragraphs[0].add_run(input_name)
    table_issue.rows[2].cells[1].paragraphs[0].add_run(input_risk)
    table_issue.rows[3].cells[1].paragraphs[0].add_run(input_type)
    table_issue.rows[4].cells[1].paragraphs[0].add_run(input_target)
    url_count = 0
    for i in input_range:
        if url_count == 0:
            table_issue.rows[5].cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            table_issue.rows[5].cells[1].paragraphs[0].add_run(i)
        else:
            table_issue.rows[5].cells[1].add_paragraph('')
            table_issue.rows[5].cells[1].paragraphs[url_count].paragraph_format.first_line_indent = Cm(0)
            table_issue.rows[5].cells[1].paragraphs[url_count].add_run(i)
        url_count += 1
        if url_count > 9:
            table_issue.rows[5].cells[1].add_paragraph('')
            table_issue.rows[5].cells[1].paragraphs[url_count].paragraph_format.first_line_indent = Cm(0)
            table_issue.rows[5].cells[1].paragraphs[url_count].add_run('⋮')
            break
    table_issue.rows[6].cells[1].paragraphs[0].add_run(input_detail)
    table_issue.rows[7].cells[1].paragraphs[0].add_run(input_advice)
    info_count = 0
    for temp_info in input_info:
        if info_count == 0:
            table_issue.rows[8].cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
            table_issue.rows[8].cells[1].paragraphs[0].add_run(temp_info)
        else:
            table_issue.rows[8].cells[1].add_paragraph('')
            table_issue.rows[8].cells[1].paragraphs[info_count].paragraph_format.first_line_indent = Cm(0)
            table_issue.rows[8].cells[1].paragraphs[info_count].add_run(temp_info)
        info_count += 1
        if info_count > 9:
            table_issue.rows[8].cells[1].add_paragraph('')
            table_issue.rows[8].cells[1].paragraphs[info_count].paragraph_format.first_line_indent = Cm(0)
            table_issue.rows[8].cells[1].paragraphs[info_count].add_run('⋮')
            break
    # table_issue.rows[8].cells[1].paragraphs[0].add_run(input_info)
    new_table = deepcopy(table_issue)
    # write table to document
    x = input_document.paragraphs
    paragraph = x[70]
    para2 = input_document.add_page_break()
    paragraph._p.addnext(para2._p)
    paragraph._p.addnext(new_table._element)
    input_document.save(input_fn)  # 保存


def transfer_report(report_data_list):
    # read input

    if default_company == '':
        input_company = '範例公司'
    else:
        input_company = default_company
    input_fn = 'VAS弱掃報告_' + input_company + '.docx'

    if default_date_start == '':
        input_date_1 = read_date_yesterday()
        input_date_2 = read_date_today()
    else:
        input_date_1 = default_date_start
        input_date_2 = default_date_before

    if default_date_end == '':
        input_date_3 = read_date_today()
    else:
        input_date_3 = default_date_end

    if default_screenshot_url == '':
        input_url = 'https://www.example.com'
    else:
        input_url = default_screenshot_url

    issue_count = 1
    issue_list = []
    count_high = 0
    count_medium = 0
    count_low = 0
    website_count = 0

    for report_data in report_data_list:
        if str(report_data['site'][0]['@name']) != '':
            website_count += 1
        try:
            # domain = read_domain(data_url)
            # read issue list
            for site in report_data['site']:
                # temp_alerts = reversed(site['alerts'])
                for alert in site['alerts']:
                    # read issue level
                    temp_level = alert['riskdesc'].split('(')
                    level = temp_level[0]
                    if "Informational" in level:
                        continue
                    elif "Low" in level:
                        level = '低'
                    elif "Medium" in level:
                        level = '中'
                    elif "High" in level:
                        level = '高'
                    # read issue advice
                    advice = alert['solution']
                    advice = advice.replace("<p>", "")
                    advice = advice.replace("</p>", "")
                    advice = remove_punctuation(advice)
                    # read issue detail
                    detail = alert['desc']
                    detail = detail.replace("<p>", "")
                    detail = detail.replace("</p>", "")
                    detail = remove_punctuation(detail)
                    # read type (cwe)
                    try:
                        cwe_id = alert['cweid']
                    except Exception as ex:
                        cwe_id = ''
                        print(str(ex))

                    temp_instances = alert['instances']
                    temp_url_list = []
                    temp_evidence_list = []
                    for temp_single in temp_instances:
                        temp_url_list.append(temp_single['uri'])
                        if not temp_single['evidence'] == '':
                            temp_evidence_list.append(temp_single['uri'] + '(' + temp_single['evidence'] + ')')
                    temp_url_list = list(set(temp_url_list))

                    # test backup
                    report_issue_name = alert['name']
                    report_issue_detail = detail
                    report_issue_advice = advice
                    report_issue_type = 'CWE-' + str(cwe_id)
                    report_status = True
                    report_issue_cost = 'Medium'
                    # test en to zh
                    try:
                        df = pandas.read_excel("zap_report.xlsx")
                        nmp = df.values
                        for n in nmp:
                            if str(n[0]) == alert['pluginid']:
                                if not n[3]:
                                    report_status = False
                                    break
                                print('id:' + str(n[0]))
                                print(report_issue_name)
                                if str(n[1]) != "nan":
                                    if str(n[1]) == 'A1':
                                        report_issue_type = "OWASP A01:2021-權限控制失效"
                                    if str(n[1]) == 'A2':
                                        report_issue_type = "OWASP A02:2021-加密機制失效"
                                    if str(n[1]) == 'A3':
                                        report_issue_type = "OWASP A03:2021-注入式攻擊"
                                    if str(n[1]) == 'A4':
                                        report_issue_type = "OWASP A04:2021-不安全設計"
                                    if str(n[1]) == 'A5':
                                        report_issue_type = "OWASP A05:2021-安全設定缺陷"
                                    if str(n[1]) == 'A6':
                                        report_issue_type = "OWASP A06:2021-危險或過舊的元件"
                                    if str(n[1]) == 'A7':
                                        report_issue_type = "OWASP A07:2021-認證及驗證機制失效"
                                    if str(n[1]) == 'A8':
                                        report_issue_type = "OWASP A08:2021-軟體及資料完整性失效"
                                    if str(n[1]) == 'A9':
                                        report_issue_type = "OWASP A09:2021-資安記錄及監控失效"
                                    if str(n[1]) == 'A10':
                                        report_issue_type = "OWASP A10:2021-伺服端請求偽造"
                                if str(n[5]) != "nan":
                                    report_issue_cost = str(n[5])
                                if str(n[7]) != "nan":
                                    report_issue_name = str(n[7])
                                if str(n[9]) != "nan":
                                    report_issue_detail = str(n[9])
                                if str(n[11]) != "nan":
                                    report_issue_advice = str(n[11])

                    except Exception as ex:
                        print(ex)
                    if report_status:
                        check_status = True
                        for check_issue in issue_list:
                            if check_issue['name'] == report_issue_name:
                                check_issue['url'] = check_issue['url'] + temp_url_list
                                check_issue['info'] = check_issue['info'] + temp_evidence_list
                                if report_data['target'] not in check_issue['target']:
                                    check_issue['target'] = check_issue['target'] + '\n' + report_data['target']
                                check_status = False
                                break
                            else:
                                pass
                        if check_status:
                            issue_dict = {'id': 'VAS' + str(issue_count).zfill(2),
                                          'name': report_issue_name,
                                          'level': level,
                                          'cost': report_issue_cost,
                                          'type': report_issue_type,
                                          'target': report_data['target'],
                                          'url': temp_url_list,
                                          'detail': report_issue_detail,
                                          'advice': report_issue_advice,
                                          'info': temp_evidence_list}
                            issue_list.append(issue_dict)
                            issue_count += 1

        except Exception as ex:
            print(ex)

    for temp_issue in issue_list:
        if temp_issue['level'] == '高':
            count_high += 1
        if temp_issue['level'] == '中':
            count_medium += 1
        if temp_issue['level'] == '低':
            count_low += 1
    # test

    custom_image.write_result(issue_list)
    # update image
    path = "template_vas_zh.docx"
    doc = DocxTemplate(path)
    custom_chart.zap_score(count_low, count_medium, count_high)
    if count_low == 0 and count_medium == 0 and count_high == 0:
        shutil.copyfile('none_issue.jpg', 'temp_distribution.jpg')
    else:
        custom_chart.zap_pie(count_low, count_medium, count_high)

    temp_test = """
    依照掃描結果可將問題分為2大類，再針對各類別進行更詳細的說明與建議。

    1. 不安全的設計：
      ◼ 建立與使用安全開發生命週期並且協同程式安全的專業人士來設計安全控制措施。
      ◼ 建立與使用安全設計模式的函式庫或是已完成可使用的元件。
      ◼ 使用威脅建模在關鍵的認證、存取控制、商業邏輯與關鍵缺陷上。
      ◼ 撰寫單元測試與整合測試來驗證所有的關鍵流程對威脅建模都有抵抗。
    2. 未進行最佳化安全設定：
      ◼ 快速且簡單的佈署，而且能在分隔且封鎖的環境下執行。
      ◼ 開發，品質管理，以及實際營運的環境，都須有一致相同的設定，並且使用不同的認證資訊。
      ◼ 不會搭配任何不需要的功能，套件，檔案，以及範本。
    """

    if website_count == 0:
        temp_test = ' 本次掃描中未發現網站服務，請確認測試目標狀態其網站服務是否有正確運行。'

    replacements = {
        'replace_company': input_company,
        'replace_date_1': input_date_1,
        'replace_date_2': input_date_2,
        'replace_date_3': input_date_3,
        'replace_target_count': len(report_data_list),
        'replace_website_count': str(website_count),
        'replace_summary': temp_test,
        'replace_url': input_url,
        'image_score': InlineImage(doc, 'temp_score.jpg', width=Mm(90)),
        'image_dis': InlineImage(doc, 'temp_distribution.jpg', width=Mm(90)),
        'image_9': InlineImage(doc, 'temp_grid.png', width=Mm(180)),
    }
    doc.render(replacements)  # 渲染替换
    doc.save(input_fn)  # 保存

    # clear image
    os.remove('temp_score.jpg')
    os.remove('temp_distribution.jpg')
    os.remove('temp_grid.png')

    for temp_issue in reversed(issue_list):
        add_issue(input_fn,
                  temp_issue['id'],
                  temp_issue['name'],
                  temp_issue['level'],
                  temp_issue['type'],
                  temp_issue['target'],
                  temp_issue['url'],
                  temp_issue['detail'],
                  temp_issue['advice'],
                  temp_issue['info'],
                  input_fn,
                  False)
    for temp_issue in issue_list:
        add_summary(input_fn, temp_issue['id'], temp_issue['name'], temp_issue['level'], input_fn, False)

    target_count = 0
    for report_data in report_data_list:
        target_count += 1
        if str(report_data['site'][0]['@name']) == '':
            add_target(input_fn, report_data['target'], str(report_data['input']))
            continue
        else:
            add_target(input_fn, report_data['target'], str(report_data['site'][0]['@name']))
        if default_screenshot_url == '':
            read_screenshot(str(report_data['site'][0]['@name']), report_data['target'] + '.png')
        else:
            read_screenshot(default_screenshot_url, report_data['target'] + '.png')
        add_screenshot(input_fn, report_data['target'] + '.png',
                       str(target_count) + '.' + report_data['target'] + '網站進入點:')

        os.remove(report_data['target'] + '.png')


def transfer_report_en(report_data_list):
    # read input
    if default_company == '':
        input_company = 'Example Company'
    else:
        input_company = default_company
    input_fn = 'VAS_Report_' + input_company + '.docx'

    if default_date_start == '':
        input_date_1 = read_date_yesterday()
        input_date_2 = read_date_today()
    else:
        input_date_1 = default_date_start
        input_date_2 = default_date_before

    if default_date_end == '':
        input_date_3 = read_date_today()
    else:
        input_date_3 = default_date_end

    if default_screenshot_url == '':
        input_url = 'https://www.example.com'
    else:
        input_url = default_screenshot_url

    issue_count = 1
    issue_list = []
    count_high = 0
    count_medium = 0
    count_low = 0
    website_count = 0

    for report_data in report_data_list:
        try:
            if str(report_data['site'][0]['@name']) != '':
                website_count += 1
            # read issue list
            for site in report_data['site']:
                # temp_alerts = reversed(site['alerts'])
                for alert in site['alerts']:
                    # read issue level
                    temp_level = alert['riskdesc'].split('(')
                    level = temp_level[0]
                    if "Informational" in level:
                        continue
                    # read issue advice
                    advice = alert['solution']
                    advice = advice.replace("<p>", "")
                    advice = advice.replace("</p>", "")
                    advice = remove_punctuation(advice)
                    # read issue detail
                    detail = alert['desc']
                    detail = detail.replace("<p>", "")
                    detail = detail.replace("</p>", "")
                    detail = remove_punctuation(detail)
                    # read type (cwe)
                    try:
                        cwe_id = alert['cweid']
                    except Exception as ex:
                        cwe_id = ''
                        utils.custom_log.debug(str(ex))

                    temp_instances = alert['instances']
                    temp_url_list = []
                    temp_evidence_list = []
                    for temp_single in temp_instances:
                        temp_url_list.append(temp_single['uri'])
                        if not temp_single['evidence'] == '':
                            temp_evidence_list.append(temp_single['uri'] + '(' + temp_single['evidence'] + ')')
                    temp_url_list = list(set(temp_url_list))

                    # test backup
                    report_issue_name = alert['name']
                    report_issue_detail = detail
                    report_issue_advice = advice
                    report_issue_type = 'CWE-' + str(cwe_id)
                    report_status = True
                    report_issue_cost = 'Medium'
                    # test en to zh
                    try:
                        df = pandas.read_excel("zap_report.xlsx")
                        nmp = df.values
                        for n in nmp:
                            if str(n[0]) == alert['pluginid']:
                                if not n[3]:
                                    report_status = False
                                    break
                                if str(n[1]) != "nan":
                                    if str(n[1]) == 'A1':
                                        report_issue_type = "OWASP A01:2021-Broken Access Control"
                                    if str(n[1]) == 'A2':
                                        report_issue_type = "OWASP A02:2021-Cryptographic Failures "
                                    if str(n[1]) == 'A3':
                                        report_issue_type = "OWASP A03:2021-Injection"
                                    if str(n[1]) == 'A4':
                                        report_issue_type = "OWASP A04:2021-Insecure Design"
                                    if str(n[1]) == 'A5':
                                        report_issue_type = "OWASP A05:2021-Security Misconfiguration"
                                    if str(n[1]) == 'A6':
                                        report_issue_type = "OWASP A06:2021-Vulnerable and Outdated Components"
                                    if str(n[1]) == 'A7':
                                        report_issue_type = "OWASP A07:2021-Identification and Authentication Failures"
                                    if str(n[1]) == 'A8':
                                        report_issue_type = "OWASP A08:2021-Software and Data Integrity Failures"
                                    if str(n[1]) == 'A9':
                                        report_issue_type = "OWASP A09:2021-Security Logging and Monitoring Failures"
                                    if str(n[1]) == 'A10':
                                        report_issue_type = "OWASP A10:2021-Server-Side Request Forgery"

                                if str(n[5]) != "nan":
                                    report_issue_cost = str(n[4])
                                if str(n[7]) != "nan":
                                    report_issue_name = str(n[6])
                                if str(n[9]) != "nan":
                                    report_issue_detail = str(n[8])
                                if str(n[11]) != "nan":
                                    report_issue_advice = str(n[10])

                    except Exception as ex:
                        print(ex)
                    if report_status:
                        check_status = True
                        for check_issue in issue_list:
                            if check_issue['name'] == report_issue_name:
                                check_issue['url'] = check_issue['url'] + temp_url_list
                                check_issue['info'] = check_issue['info'] + temp_evidence_list
                                if report_data['target'] not in check_issue['target']:
                                    check_issue['target'] = check_issue['target'] + '\n' + report_data['target']
                                check_status = False
                                break
                            else:
                                pass
                        if check_status:
                            issue_dict = {'id': 'VAS' + str(issue_count).zfill(2),
                                          'name': report_issue_name,
                                          'level': level,
                                          'cost': report_issue_cost,
                                          'type': report_issue_type,
                                          'target': report_data['target'],
                                          'url': temp_url_list,
                                          'detail': report_issue_detail,
                                          'advice': report_issue_advice,
                                          'info': temp_evidence_list}
                            issue_list.append(issue_dict)
                            print("id:" + issue_dict['id'])
                            issue_count += 1

            for temp_issue in issue_list:
                if temp_issue['level'] == 'High ':
                    count_high += 1
                if temp_issue['level'] == 'Medium ':
                    count_medium += 1
                if temp_issue['level'] == 'Low ':
                    count_low += 1

            # screenshot
        except Exception as ex:
            print(ex)

    custom_image.write_result_en(issue_list)

    # update image
    path = "template_vas_en.docx"
    doc = DocxTemplate(path)
    custom_chart.zap_score(count_low, count_medium, count_high)
    if count_low == 0 and count_medium == 0 and count_high == 0:
        shutil.copyfile('none_issue.jpg', 'temp_distribution.jpg')
    else:
        custom_chart.zap_pie_en(count_low, count_medium, count_high)

    temp_test = """
    According to the scan results, the problems can be divided into two categories, and then more detailed 
    explanations and suggestions are made for each category

    1. Insecure Design：
    ◼ Establish and use the security development lifecycle and collaborate with application security
    professionals to evaluate controls related to design security and privacy.	
    ◼ Establish and use a library of safe design patterns or complete a component that can be used.
    ◼ Use threat modeling on critical authentication, access control, business logic, and critical flaws.
    ◼ Write unit tests and integration tests to verify that all critical processes are resistant to threat modeling.

    2. Security Misconfiguration：
    ◼ A repeatable security hardening process that can be deployed quickly and easily and can be executed in
    a compartmentalized and blocked environment. Development, quality management, and actual operation 
    environments must all have the same settings and use different certification information. This step needs to
    be automated as much as possible, reducing the investment required to set up a secure environment.
    ◼ A minimal platform that doesn't come with any unwanted features, suites, archives, and templates. 
    Remove or do not install any features or frameworks that you do not need to use.
"""

    if website_count == 0:
        temp_test = """
        No website service was found in this scan. Please make sure the website is available."""

    replacements = {
        'replace_company': input_company,
        'replace_date_1': input_date_1,
        'replace_date_2': input_date_2,
        'replace_date_3': input_date_3,
        'replace_target_count': len(report_data_list),
        'replace_website_count': str(website_count),
        'replace_summary': temp_test,
        'replace_url': input_url,
        'image_score': InlineImage(doc, 'temp_score.jpg', width=Mm(90)),
        'image_dis': InlineImage(doc, 'temp_distribution.jpg', width=Mm(90)),
        'image_9': InlineImage(doc, 'temp_grid.png', width=Mm(180)),
    }

    doc.render(replacements)  # 渲染替换
    doc.save(input_fn)  # 保存

    # clear image
    os.remove('temp_score.jpg')
    os.remove('temp_distribution.jpg')
    os.remove('temp_grid.png')

    for temp_issue in reversed(issue_list):
        add_issue(input_fn,
                  temp_issue['id'],
                  temp_issue['name'],
                  temp_issue['level'],
                  temp_issue['type'],
                  temp_issue['target'],
                  temp_issue['url'],
                  temp_issue['detail'],
                  temp_issue['advice'],
                  temp_issue['info'],
                  input_fn,
                  True)
    for temp_issue in issue_list:
        add_summary(input_fn, temp_issue['id'], temp_issue['name'], temp_issue['level'], input_fn, True)

    target_count = 1
    for report_data in report_data_list:
        if str(report_data['site'][0]['@name']) == '':
            add_target(input_fn, report_data['target'], str(report_data['input']))
            continue
        else:
            add_target(input_fn, report_data['target'], str(report_data['site'][0]['@name']))

        if default_screenshot_url == '':
            read_screenshot(str(report_data['site'][0]['@name']), report_data['target'] + '.png')
        else:
            read_screenshot(default_screenshot_url, report_data['target'] + '.png')
        add_screenshot(input_fn, report_data['target'] + '.png',
                       str(target_count) + '.' + report_data['target'] + ' access url:')
        target_count += 1
