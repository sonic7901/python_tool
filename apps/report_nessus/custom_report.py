import os
import json
import re
import shutil
import pandas
import logging
from datetime import datetime, timedelta
from copy import deepcopy
from docx import Document
from docx.shared import Cm, Mm
from docxtpl import DocxTemplate, InlineImage
import custom_image
import custom_chart
import custom_request
import custom_csv

# manual input
default_company = ''
default_date_before = ''
default_date_start = ''
default_date_end = ''


def set_company(input_company):
    try:
        global default_company
        default_company = input_company
    except Exception as ex:
        logging.error('Exception:' + str(ex))


def set_date_start(input_date_start):
    try:
        global default_date_start
        global default_date_before
        datetime_object = datetime.strptime(input_date_start, "%Y/%m/%d")
        default_date_start = datetime_object.strftime("%Y/%m/%d")
        default_date_before = (datetime_object + timedelta(1)).strftime("%Y/%m/%d")
    except Exception as ex:
        logging.error('Exception:' + str(ex))


def set_date_end(input_date_end):
    try:
        global default_date_end
        default_date_end = input_date_end
    except Exception as ex:
        logging.error('Exception:' + str(ex))


def read_json(report_name):
    dict_data = ""
    try:
        if os.path.isfile(report_name):
            with open(report_name, 'r', encoding='utf-8') as read_file:
                dict_data = json.load(read_file)
    except Exception as ex:
        print('Exception:' + str(ex))
    return dict_data


def read_date_today():
    try:
        time_now = datetime.now()
        str_now = time_now.strftime("%Y/%m/%d")
    except Exception as ex:
        logging.error('Exception:' + str(ex))
    return str_now


def read_date_yesterday():
    try:
        time_yesterday = datetime.now() - timedelta(1)
        str_yesterday = time_yesterday.strftime("%Y/%m/%d")
        return str_yesterday
    except Exception as ex:
        logging.error('Exception:' + str(ex))


def write_variable(temp_document, replace_dict, input_fn):
    try:
        doc = DocxTemplate(temp_document)  # 选定模板
        doc.render(replace_dict)  # 渲染替换
        doc.save(input_fn)  # 保存
    except Exception as ex:
        logging.error('Exception:' + str(ex))


def remove_punctuation(text):
    try:
        punctuation = '!,;:?"\''
        text = re.sub(r'[{}]+'.format(punctuation), '', text)
        return text.strip().lower()
    except Exception as ex:
        print('Exception' + str(ex))
        return ''


def add_summary(input_docx, input_id, input_name, input_risk, input_filename, input_en):
    try:
        input_document = Document(input_docx)
        tables = input_document.tables
        table = tables[1]
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].paragraph_format.first_line_indent = Cm(0.45)
        row_cells[0].paragraphs[0].add_run(input_id)
        row_cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        row_cells[1].paragraphs[0].add_run(input_name)
        if input_en:
            row_cells[2].paragraphs[0].paragraph_format.first_line_indent = Cm(0.45)
            row_cells[2].paragraphs[0].add_run(input_risk)
        else:
            row_cells[2].text = input_risk
        input_document.save(input_filename)
    except Exception as ex:
        logging.error('Exception:' + str(ex))
    return


def add_target(input_docx, input_name, input_url):
    try:
        input_document = Document(input_docx)
        tables = input_document.tables
        table = tables[0]
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        row_cells[0].paragraphs[0].add_run(input_name)
        row_cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        row_cells[1].paragraphs[0].add_run(input_url)
        input_document.save(input_docx)
    except Exception as ex:
        logging.error('Exception:' + str(ex))
    return


def add_issue(input_docx,
              input_id,
              input_name,
              input_risk,
              input_type,
              input_range,
              input_detail,
              input_advice,
              input_info,
              input_fn,
              input_en):
    try:
        # create issue table
        input_document = Document(input_docx)
        if input_en:
            document_table = Document('template_table_issue_en.docx')
        else:
            document_table = Document('template_table_issue_zh.docx')
        temp_table = document_table.tables
        table_issue = temp_table[0]
        table_issue.rows[0].cells[1].paragraphs[0].add_run(input_id)
        table_issue.rows[1].cells[1].paragraphs[0].add_run(input_name)
        table_issue.rows[2].cells[1].paragraphs[0].add_run(input_risk)
        table_issue.rows[3].cells[1].paragraphs[0].add_run(input_type)
        url_count = 0
        for i in input_range:
            if url_count == 0:
                table_issue.rows[4].cells[1].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
                table_issue.rows[4].cells[1].paragraphs[0].add_run(i)
            else:
                table_issue.rows[4].cells[1].add_paragraph('')
                table_issue.rows[4].cells[1].paragraphs[url_count].paragraph_format.first_line_indent = Cm(0)
                table_issue.rows[4].cells[1].paragraphs[url_count].add_run(i)
            url_count += 1
            if url_count > 9:
                table_issue.rows[4].cells[1].add_paragraph('')
                table_issue.rows[4].cells[1].paragraphs[url_count].paragraph_format.first_line_indent = Cm(0)
                table_issue.rows[4].cells[1].paragraphs[url_count].add_run('⋮')
                break
        table_issue.rows[5].cells[1].paragraphs[0].add_run(input_detail)
        table_issue.rows[6].cells[1].paragraphs[0].add_run(input_advice)
        table_issue.rows[7].cells[1].paragraphs[0].add_run(input_info)
        new_table = deepcopy(table_issue)
        # write table to document
        x = input_document.paragraphs
        paragraph = x[70]
        para2 = input_document.add_page_break()
        paragraph._p.addnext(para2._p)
        paragraph._p.addnext(new_table._element)
        input_document.save(input_fn)  # 保存
    except Exception as ex:
        logging.error('Exception:' + str(ex))


def transfer_report():
    try:
        # 1. read manual input
        if default_company == '':
            input_company = '範例公司'
        else:
            input_company = default_company
        input_fn = 'VAS弱掃報告_' + input_company + '.docx'

        if default_date_start == '':
            input_date_1 = read_date_yesterday()
            input_date_2 = read_date_today()
        else:
            input_date_1 = default_date_start
            input_date_2 = default_date_before

        if default_date_end == '':
            input_date_3 = read_date_today()
        else:
            input_date_3 = default_date_end

        issue_count = 1
        issue_list = []
        count_high = 0
        count_medium = 0
        count_low = 0
        list_report_name = []
        list_report_level = []
        list_report_type = []
        list_report_range = []
        list_report_detail = []
        list_report_advice = []
        list_report_info = []
        list_report_plugin = []
        list_report_ip = []
        # 2. read critical issue to csv
        for filename in os.listdir():
            if filename.endswith('.csv'):
                print("read csv: " + filename)
                try:
                    temp_dict_list = custom_csv.read_file_to_dict(filename)
                    for temp_dict in temp_dict_list:
                        if not temp_dict["Host"] in list_report_ip:
                            list_report_ip.append(temp_dict["Host"])
                        if temp_dict["Risk"] == 'Critical':
                            temp_dict["Risk"] = '高'
                            if temp_dict["Plugin ID"] in list_report_plugin:
                                temp_num = list_report_plugin.index(temp_dict["Plugin ID"])
                                if temp_dict["Host"] not in list_report_range[temp_num]:
                                    list_report_range[temp_num].append(temp_dict["Host"])
                            else:
                                list_report_name.append(temp_dict["Name"])
                                list_report_level.append(temp_dict["Risk"])
                                list_report_type.append(temp_dict["Plugin ID"])
                                list_report_range.append([temp_dict["Host"]])
                                list_report_detail.append(temp_dict["Description"])
                                list_report_advice.append(temp_dict["Solution"])
                                list_report_info.append(temp_dict["See Also"])
                                list_report_plugin.append(temp_dict["Plugin ID"])
                except Exception as ex:
                    print('Exception(read csv):' + str(ex))
        # 3. read high issue to csv
        for filename in os.listdir():
            if filename.endswith('.csv'):
                print("read csv: " + filename)
                temp_dict_list = custom_csv.read_file_to_dict(filename)
                try:
                    # High
                    for temp_dict in temp_dict_list:
                        if not temp_dict["Host"] in list_report_ip:
                            list_report_ip.append(temp_dict["Host"])
                        if temp_dict["Risk"] == 'High':
                            temp_dict["Risk"] = '高'
                            if temp_dict["Plugin ID"] in list_report_plugin:
                                temp_num = list_report_plugin.index(temp_dict["Plugin ID"])
                                if temp_dict["Host"] not in list_report_range[temp_num]:
                                    list_report_range[temp_num].append(temp_dict["Host"])
                            else:
                                list_report_name.append(temp_dict["Name"])
                                list_report_level.append(temp_dict["Risk"])
                                list_report_type.append(temp_dict["Plugin ID"])
                                list_report_range.append([temp_dict["Host"]])
                                list_report_detail.append(temp_dict["Description"])
                                list_report_advice.append(temp_dict["Solution"])
                                list_report_info.append(temp_dict["See Also"])
                                list_report_plugin.append(temp_dict["Plugin ID"])
                except Exception as ex:
                    print('Exception:' + str(ex))
        # 4. read medium issue to csv
        for filename in os.listdir():
            if filename.endswith('.csv'):
                print("read csv: " + filename)
                temp_dict_list = custom_csv.read_file_to_dict(filename)
                try:
                    # Medium
                    for temp_dict in temp_dict_list:
                        if not temp_dict["Host"] in list_report_ip:
                            list_report_ip.append(temp_dict["Host"])
                        if temp_dict["Risk"] == 'Medium':
                            temp_dict["Risk"] = '中'
                            if temp_dict["Plugin ID"] in list_report_plugin:
                                temp_num = list_report_plugin.index(temp_dict["Plugin ID"])
                                if temp_dict["Host"] not in list_report_range[temp_num]:
                                    list_report_range[temp_num].append(temp_dict["Host"])
                            else:
                                list_report_name.append(temp_dict["Name"])
                                list_report_level.append(temp_dict["Risk"])
                                list_report_type.append(temp_dict["Plugin ID"])
                                list_report_range.append([temp_dict["Host"]])
                                list_report_detail.append(temp_dict["Description"])
                                list_report_advice.append(temp_dict["Solution"])
                                list_report_info.append(temp_dict["See Also"])
                                list_report_plugin.append(temp_dict["Plugin ID"])
                except Exception as ex:
                    print('Exception:' + str(ex))
        # 4. read low issue to csv
        for filename in os.listdir():
            if filename.endswith('.csv'):
                print("read csv: " + filename)
                temp_dict_list = custom_csv.read_file_to_dict(filename)
                try:
                    # Low
                    for temp_dict in temp_dict_list:
                        if not temp_dict["Host"] in list_report_ip:
                            list_report_ip.append(temp_dict["Host"])
                        if temp_dict["Risk"] == 'Low':
                            temp_dict["Risk"] = '低'
                            if temp_dict["Plugin ID"] in list_report_plugin:
                                temp_num = list_report_plugin.index(temp_dict["Plugin ID"])
                                if temp_dict["Host"] not in list_report_range[temp_num]:
                                    list_report_range[temp_num].append(temp_dict["Host"])
                            else:
                                list_report_name.append(temp_dict["Name"])
                                list_report_level.append(temp_dict["Risk"])
                                list_report_type.append(temp_dict["Plugin ID"])
                                list_report_range.append([temp_dict["Host"]])
                                list_report_detail.append(temp_dict["Description"])
                                list_report_advice.append(temp_dict["Solution"])
                                list_report_info.append(temp_dict["See Also"])
                                list_report_plugin.append(temp_dict["Plugin ID"])
                except Exception as ex:
                    print('Exception:' + str(ex))
        # 5. en to zh
        print("issue number: " + str(len(list_report_name)))
        for i in range(len(list_report_name)):
            temp_zh_result = custom_request.read_nessus(list_report_type[i])
            print(str(i + 1) + ". " + temp_zh_result[0])
            issue_dict = {'id': 'VAS' + str(i + 1).zfill(2),
                          'name': temp_zh_result[0],
                          'level': list_report_level[i],
                          'type': temp_zh_result[3],
                          'range': list_report_range[i],
                          'detail': temp_zh_result[1],
                          'advice': temp_zh_result[2],
                          'info': list_report_info[i]}
            issue_count += 1
            issue_list.append(issue_dict)
        # 6. score
        for temp_issue in issue_list:
            if temp_issue['level'] == '高':
                count_high += 1
            if temp_issue['level'] == '中':
                count_medium += 1
            if temp_issue['level'] == '低':
                count_low += 1
        # 7. write word report
        path = "template_vas_zh.docx"
        doc = DocxTemplate(path)
        custom_chart.zap_score(count_low, count_medium, count_high)
        if count_low == 0 and count_medium == 0 and count_high == 0:
            shutil.copyfile('none_issue.jpg', 'temp_distribution.jpg')
        else:
            custom_chart.zap_pie(count_low, count_medium, count_high)
        temp_zh_summary = """
    依照掃描結果可將問題分為 2 大類，再針對各類別進行更詳細的說明與建議。

    1. 使用了未更新或不安全的服務：
      ◼ 安裝最新的安全更新和系統更新。
      定期更新系統和軟體以獲得最新的修補程式和功能改進，確保主機不受已知漏洞的影響。
      ◼ 限制不必要的網路服務和連接。
      網絡中的服務和連接可以成為攻擊的入口，建議限制主機上不必要服務，以降低被攻擊的風險。
      ◼ 使用具加密協議的服務，如 HTTPS 和 SFTP。
      ◼ 禁用含不安全的協議的服務，如 FTP 和 Telnet。
    2. 未進行最佳化安全設定：
      ◼ 關閉未使用的遠程管理。
      禁止遠端管理可以有效減少攻擊者能選擇的攻擊方式，降低被攻擊的風險。
      ◼ 加強文件和目錄的訪問控制。
      增强文件和目錄的訪問控制可以減少資訊外洩的機會，系統資訊洩漏有機會衍伸成其他種類攻擊。
      ◼ 設定防火牆規則。
      在內網設置適當的防火牆規則可以防止攻擊在內部橫向移動，也可以阻止不安全的連接進入網絡。
        """
        website_count = len(list_report_ip)
        if website_count == 0:
            temp_zh_summary = ' 本次掃描中未發現網路服務，請確認測試目標是否有正確運行。'
        target_count = 0
        target_list = []
        with open("ip.txt", "r") as file:
            for line in file:
                if not line == "":
                    target_list.append(line.strip())
                    target_count += 1
        for temp_ip in target_list:
            if temp_ip not in list_report_ip:
                logging.info("host not found in result: " + temp_ip)
        for temp_ip in list_report_ip:
            if temp_ip not in target_list:
                logging.info("host not found in target: " + temp_ip)
        replacements = {
            'replace_company': input_company,
            'replace_date_1': input_date_1,
            'replace_date_2': input_date_2,
            'replace_date_3': input_date_3,
            'replace_target_count': str(target_count),
            'replace_website_count': str(website_count),
            'replace_summary': temp_zh_summary,
            'replace_url': "input_url",
            'image_score': InlineImage(doc, 'temp_score.jpg', width=Mm(90)),
            'image_dis': InlineImage(doc, 'temp_distribution.jpg', width=Mm(90)),
        }
        doc.render(replacements)
        doc.save(input_fn)
        os.remove('temp_score.jpg')
        os.remove('temp_distribution.jpg')
        for temp_issue in reversed(issue_list):
            add_issue(input_fn,
                      temp_issue['id'],
                      temp_issue['name'],
                      temp_issue['level'],
                      temp_issue['type'],
                      temp_issue['range'],
                      temp_issue['detail'],
                      temp_issue['advice'],
                      temp_issue['info'],
                      input_fn,
                      False)
        for temp_issue in issue_list:
            add_summary(input_fn, temp_issue['id'], temp_issue['name'], temp_issue['level'], input_fn, False)
        device_count = 0
        target_list.sort()
        for temp_target in target_list:
            device_count += 1
            add_target(input_fn, "Device_" + str(device_count), temp_target.strip())
    except Exception as ex:
        logging.error('Exception:' + str(ex))


def transfer_report_en(report_data_list):
    # read input
    if default_company == '':
        input_company = 'Example Company'
    else:
        input_company = default_company
    input_fn = 'VAS_Report_' + input_company + '.docx'

    if default_date_start == '':
        input_date_1 = read_date_yesterday()
        input_date_2 = read_date_today()
    else:
        input_date_1 = default_date_start
        input_date_2 = default_date_before

    if default_date_end == '':
        input_date_3 = read_date_today()
    else:
        input_date_3 = default_date_end

    issue_count = 1
    issue_list = []
    count_high = 0
    count_medium = 0
    count_low = 0
    website_count = 0

    for report_data in report_data_list:
        try:
            if str(report_data['site'][0]['@name']) != '':
                website_count += 1
            # read issue list
            for site in report_data['site']:
                # temp_alerts = reversed(site['alerts'])
                for alert in site['alerts']:
                    # read issue level
                    temp_level = alert['riskdesc'].split('(')
                    level = temp_level[0]
                    if alert['pluginid'] == '2':
                        continue
                    if alert['pluginid'] == '3':
                        continue
                    if '10020' in alert['pluginid']:
                        continue
                    if "Informational" in level:
                        continue
                    # read issue advice
                    advice = alert['solution']
                    advice = advice.replace("<p>", "")
                    advice = advice.replace("</p>", "")
                    advice = remove_punctuation(advice)
                    # read issue detail
                    detail = alert['desc']
                    detail = detail.replace("<p>", "")
                    detail = detail.replace("</p>", "")
                    detail = remove_punctuation(detail)
                    # read type (cwe)
                    try:
                        cwe_id = alert['cweid']
                    except Exception as ex:
                        cwe_id = ''
                        print(str(ex))

                    temp_instances = alert['instances']
                    temp_url_list = []
                    temp_evidence_list = []
                    for temp_single in temp_instances:
                        temp_url_list.append(temp_single['uri'])
                        if not temp_single['evidence'] == '':
                            temp_evidence_list.append(temp_single['uri'] + '(' + temp_single['evidence'] + ')')
                    temp_url_list = list(set(temp_url_list))

                    # test backup
                    report_issue_name = alert['name']
                    report_issue_detail = detail
                    report_issue_advice = advice
                    report_issue_type = 'CWE-' + str(cwe_id)
                    report_status = True
                    report_issue_cost = 'Medium'
                    # test en to zh
                    try:
                        df = pandas.read_excel("zap_report.xlsx")
                        nmp = df.values
                        for n in nmp:
                            if str(n[0]) == alert['pluginid']:
                                if not n[3]:
                                    report_status = False
                                    break
                                if str(n[1]) != "nan":
                                    if str(n[1]) == 'A1':
                                        report_issue_type = "OWASP A01:2021-Broken Access Control"
                                    if str(n[1]) == 'A2':
                                        report_issue_type = "OWASP A02:2021-Cryptographic Failures "
                                    if str(n[1]) == 'A3':
                                        report_issue_type = "OWASP A03:2021-Injection"
                                    if str(n[1]) == 'A4':
                                        report_issue_type = "OWASP A04:2021-Insecure Design"
                                    if str(n[1]) == 'A5':
                                        report_issue_type = "OWASP A05:2021-Security Misconfiguration"
                                    if str(n[1]) == 'A6':
                                        report_issue_type = "OWASP A06:2021-Vulnerable and Outdated Components"
                                    if str(n[1]) == 'A7':
                                        report_issue_type = "OWASP A07:2021-Identification and Authentication Failures"
                                    if str(n[1]) == 'A8':
                                        report_issue_type = "OWASP A08:2021-Software and Data Integrity Failures"
                                    if str(n[1]) == 'A9':
                                        report_issue_type = "OWASP A09:2021-Security Logging and Monitoring Failures"
                                    if str(n[1]) == 'A10':
                                        report_issue_type = "OWASP A10:2021-Server-Side Request Forgery"

                                if str(n[5]) != "nan":
                                    report_issue_cost = str(n[4])
                                if str(n[7]) != "nan":
                                    report_issue_name = str(n[6])
                                if str(n[9]) != "nan":
                                    report_issue_detail = str(n[8])
                                if str(n[11]) != "nan":
                                    report_issue_advice = str(n[10])

                    except Exception as ex:
                        print(ex)
                    if report_status:
                        check_status = True
                        for check_issue in issue_list:
                            if check_issue['name'] == report_issue_name:
                                check_issue['url'] = check_issue['url'] + temp_url_list
                                check_issue['info'] = check_issue['info'] + temp_evidence_list
                                if report_data['target'] not in check_issue['target']:
                                    check_issue['target'] = check_issue['target'] + '\n' + report_data['target']
                                check_status = False
                                break
                            else:
                                pass
                        if check_status:
                            issue_dict = {'id': 'VAS' + str(issue_count).zfill(2),
                                          'name': report_issue_name,
                                          'level': level,
                                          'cost': report_issue_cost,
                                          'type': report_issue_type,
                                          'target': report_data['target'],
                                          'url': temp_url_list,
                                          'detail': report_issue_detail,
                                          'advice': report_issue_advice,
                                          'info': temp_evidence_list}
                            issue_list.append(issue_dict)
                            print("id:" + issue_dict['id'])
                            issue_count += 1

            for temp_issue in issue_list:
                if temp_issue['level'] == 'High ':
                    count_high += 1
                if temp_issue['level'] == 'Medium ':
                    count_medium += 1
                if temp_issue['level'] == 'Low ':
                    count_low += 1

            # screenshot
        except Exception as ex:
            print(ex)

    custom_image.write_result_en(issue_list)

    # update image
    path = "template_vas_en.docx"
    doc = DocxTemplate(path)
    custom_chart.zap_score(count_low, count_medium, count_high)
    if count_low == 0 and count_medium == 0 and count_high == 0:
        shutil.copyfile('none_issue.jpg', 'temp_distribution.jpg')
    else:
        custom_chart.zap_pie_en(count_low, count_medium, count_high)

    temp_test = """
    According to the scan results, the problems can be divided into two categories, and then more detailed 
    explanations and suggestions are made for each category

    1. Insecure Design：
    ◼ Establish and use the security development lifecycle and collaborate with application security
    professionals to evaluate controls related to design security and privacy.	
    ◼ Establish and use a library of safe design patterns or complete a component that can be used.
    ◼ Use threat modeling on critical authentication, access control, business logic, and critical flaws.
    ◼ Write unit tests and integration tests to verify that all critical processes are resistant to threat modeling.

    2. Security Misconfiguration：
    ◼ A repeatable security hardening process that can be deployed quickly and easily and can be executed in
    a compartmentalized and blocked environment. Development, quality management, and actual operation 
    environments must all have the same settings and use different certification information. This step needs to
    be automated as much as possible, reducing the investment required to set up a secure environment.
    ◼ A minimal platform that doesn't come with any unwanted features, suites, archives, and templates. 
    Remove or do not install any features or frameworks that you do not need to use.
"""

    if website_count == 0:
        temp_test = """
        No website service was found in this scan. Please make sure the website is available."""

    replacements = {
        'replace_company': input_company,
        'replace_date_1': input_date_1,
        'replace_date_2': input_date_2,
        'replace_date_3': input_date_3,
        'replace_target_count': len(report_data_list),
        'replace_website_count': str(website_count),
        'replace_summary': temp_test,
        'image_score': InlineImage(doc, 'temp_score.jpg', width=Mm(90)),
        'image_dis': InlineImage(doc, 'temp_distribution.jpg', width=Mm(90)),
        'image_9': InlineImage(doc, 'temp_grid.png', width=Mm(180)),
    }

    doc.render(replacements)  # 渲染替换
    doc.save(input_fn)  # 保存

    # clear image
    os.remove('temp_score.jpg')
    os.remove('temp_distribution.jpg')
    os.remove('temp_grid.png')

    for temp_issue in reversed(issue_list):
        add_issue(input_fn,
                  temp_issue['id'],
                  temp_issue['name'],
                  temp_issue['level'],
                  temp_issue['type'],
                  temp_issue['target'],
                  temp_issue['detail'],
                  temp_issue['advice'],
                  temp_issue['info'],
                  input_fn,
                  True)
    for temp_issue in issue_list:
        add_summary(input_fn, temp_issue['id'], temp_issue['name'], temp_issue['level'], input_fn, True)

    target_count = 1
    for report_data in report_data_list:
        if str(report_data['site'][0]['@name']) == '':
            add_target(input_fn, report_data['target'], str(report_data['input']))
            continue
        else:
            add_target(input_fn, report_data['target'], str(report_data['site'][0]['@name']))
        target_count += 1


if __name__ == '__main__':
    transfer_report()
